#!/usr/bin/env python3
"""Generate editable Word (.docx) report for GIC 2026 Phase 3.

Matches the scalability-first structure of generate_submission_pdf.py.
Includes GIC cover page template.
Outputs proposals/Ryoushi_Quantum_Buddies_Phase3_Report.docx.
"""
from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "results" / "phase3_final" / "figures"
OUT_DOC = ROOT / "submission" / "Quantum-Buddies_Phase3_Write-Up.docx"


def load_json(path: Path) -> dict | None:
    if not path.exists():
        return None
    with open(path) as f:
        return json.load(f)


def set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 51, 102)
    return p


def add_para(doc, text: str, bold: str = "", italic: bool = False, align: str = "left", size: int = 11) -> None:
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold:
        run = p.add_run(bold)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        p.add_run(text)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_image(doc, path: Path, caption: str, width: float = 5.5) -> None:
    if not path.exists():
        add_para(doc, f"[Figure missing: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.name = "Times New Roman"
    cap.paragraph_format.space_after = Pt(8)


def add_mermaid(doc, title: str, code: str) -> None:
    add_para(doc, title, bold="Figure (Mermaid): ", size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Inches(w)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(9)
        set_cell_shading(hdr_cells[i], "D9E2F3")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_cover_page(doc) -> None:
    add_para(doc, "Global Industry Challenge 2026 — Phase 3 Submission", bold="", size=14, align="center")
    add_para(doc, "Quantum Materials Discovery Challenge: Scaling Generative Quantum Eigensolver (GQE) Using NVIDIA CUDA-Q", bold="", size=12, align="center")
    doc.add_paragraph()
    add_para(doc, "Phase #: Phase 3", bold="", size=11, align="center")
    add_para(doc, "Team Name: Ryoushi | Quantum Buddies", bold="", size=11, align="center")
    doc.add_paragraph()

    add_para(doc, "Team Members", bold="", size=11, align="center")
    headers = ["First Name", "Last Name", "Email", "Aqora Username", "Role"]
    rows = [
        ["Gyanateet", "Dutta", "gyanateet@gmail.com", "Ryukijano", "Coder/Technical Lead"],
        ["Dat Chi (Ryan)", "Le", "ryancoltrane2004@gmail.com", "ryancdle", "Domain Expert"],
        ["Sid", "Iliyasu", "sidMelias@gmail.com", "SuperPenguin", "Business/Project Manager"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_para(doc,
             "Disclaimer: Submission must follow GIC requirements: maximum 5 pages (excluding this cover page and references), "
             "11-point Times New Roman, single spacing, and submitted via zipped folder. "
             "This cover page follows the official GIC template.",
             italic=True, size=9, align="left")
    doc.add_page_break()


def build_document() -> None:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(6)

    sections = doc.sections[0]
    sections.top_margin = Inches(1.0)
    sections.bottom_margin = Inches(1.0)
    sections.left_margin = Inches(1.0)
    sections.right_margin = Inches(1.0)

    consolidated = load_json(ROOT / "results" / "phase3_final" / "consolidated_results_gic2026.json")
    consolidated_phase3 = load_json(ROOT / "results" / "phase3_final" / "consolidated_phase3_results.json")
    ablation = load_json(ROOT / "results" / "phase3_final" / "ablation_sft_vs_rl.json")
    noise = load_json(ROOT / "results" / "phase3_final" / "noise_mitigation_summary.json")

    # ===== COVER PAGE =====
    add_cover_page(doc)

    # ===== PAGE 1: Title + Abstract + Introduction + Architecture =====
    add_para(doc, "H-cGQE: Hierarchical Conditional Generative Quantum Eigensolver", bold="", size=14, align="center")
    add_para(doc, "Quantum-Buddies Team | GIC 2026 Phase 3 Submission", align="center")
    add_para(doc, "Mitsubishi Chemical & AIST Challenge Track", align="center")
    doc.add_paragraph()

    add_heading(doc, "Abstract", level=1)
    add_para(doc,
        "We present the Hierarchical Conditional Generative Quantum Eigensolver (H-cGQE), a "
        "transformer-based architecture that amortizes quantum circuit synthesis across molecular "
        "systems via Hamiltonian-conditioned generation. H-cGQE scales from 4 to 40 qubits by "
        "combining autoregressive Pauli operator generation with MPS tensor network backends, "
        "qubit-wise commuting (QWC) measurement grouping that reduces circuit count 2-3.5x, and "
        "Sample-based Quantum Diagonalization (SQD) for noise-resilient energy extraction. On GPU, "
        "the model achieves chemical accuracy (1.6 mHa) on methyl iodide (CH3I, 0.629 mHa) and "
        "exact FCI on H2. On Rigetti Cepheus-1-108Q, we validate 12 molecules (8-28 qubits) "
        "including 8 EUV photoresist compounds relevant to Mitsubishi Chemical, with best QPU "
        "error of 13.9 mHa. DAPO reinforcement learning with a MAP-Elites archive directly "
        "optimizes for ground-state energy. We identify diagonal sequence collapse as the "
        "principal bottleneck for strongly correlated systems and propose targeted mitigations.",
        align="justify")

    add_heading(doc, "1. Introduction", level=1)
    add_para(doc,
        "The Generative Quantum Eigensolver (GQE) replaces variational optimization with "
        "autoregressive circuit generation: a transformer model proposes Pauli operator sequences "
        "that define a quantum circuit ansatz in a single forward pass, eliminating the "
        "measurement-gradient bottleneck of VQE. However, the original GQE operates on a single "
        "molecule at a time and does not address scaling beyond exact statevector simulation. "
        "Our H-cGQE extends GQE along three axes: (1) a hierarchical conditioning architecture "
        "that amortizes ansatz discovery across a molecular family, (2) QWC measurement grouping "
        "and MPS backends that extend reach from 24 to 40 qubits, and (3) a hybrid HPC-to-QPU "
        "workflow with SQD post-processing for noise-resilient energy extraction on real "
        "hardware. For Phase 3, we target EUV photoresist molecules central to Mitsubishi "
        "Chemical's industrial interests: methyl iodide, iodobenzene, phenol, o-cresol, anisole, "
        "toluene, and benzene.",
        align="justify")

    add_heading(doc, "2. Architecture", level=1)
    add_para(doc,
        "H-cGQE comprises three stages. Stage 1 (Circuit Synthesis): a Hamiltonian Encoder "
        "embeds molecular Hamiltonian terms (Pauli words + coefficients) into a dense conditioning "
        "vector via a Chemistry GNN. A GPT-2 style Operator Pool Decoder autoregressively "
        "generates Pauli operator sequences conditioned on this encoding, using UCCSD fermionic "
        "excitations mapped through Jordan-Wigner. Stage 2 (Classical Optimization): L-BFGS-B "
        "optimizes rotation angles (thetas) on CUDA-Q statevector simulators across 3x NVIDIA "
        "L40S GPUs via the nvidia-mqpu target. Stage 3 (QPU Validation): QWC Pauli terms are "
        "grouped into shared measurement circuits (2-3.5x reduction), exported as a QWC manifest, "
        "and submitted to Rigetti Cepheus-1-108Q via qBraid. SQD post-processes bitstrings with "
        "particle-number and spin-parity filtering. Training uses supervised fine-tuning on "
        "GQE-generated sequences, followed by DAPO reinforcement learning with energy-based "
        "rewards and a MAP-Elites quality-diversity archive.",
        align="justify")

    # ===== PAGE 2: Scalability (Primary Criterion) =====
    doc.add_page_break()
    add_heading(doc, "3. Scalability: From 4 to 40 Qubits", level=1)

    add_para(doc,
        "Scalability is the central challenge for quantum eigensolvers: exact statevector "
        "simulation is capped at ~24 qubits on L40S GPUs (cuStateVec distributed mode "
        "segfaults on PCIe-only systems), and VQE gradient measurements scale exponentially "
        "with system size. H-cGQE addresses this through three mechanisms:",
        align="justify")

    add_para(doc,
        "(1) QSCI + MPS Backend: Quantum-Selected Configuration Interaction samples "
        "computational-basis determinants from a quantum state, builds the Hamiltonian in "
        "that subspace, and diagonalizes classically. CUDA-Q's tensornet-mps backend extends "
        "beyond the 24q statevector cap to 40 qubits using matrix product states with "
        "controllable bond dimension (D=32,64,128,256).",
        align="justify")

    add_para(doc,
        "(2) QWC Measurement Grouping: Qubit-wise commuting Pauli terms are grouped into "
        "shared measurement circuits, reducing the number of QPU circuits by 2-3.5x. "
        "For H2 (15 terms -> 5 circuits) and LiH (631 terms -> 180 circuits), this brings "
        "both within qBraid's 2000-circuit batch limit.",
        align="justify")

    add_para(doc,
        "(3) Conditional Amortization: A single trained model generates circuits for any "
        "molecule in its training distribution, eliminating per-molecule ansatz design. "
        "The SMILES encoder provides structural priors across 10 molecules spanning 4-56 "
        "qubits, enabling the model to condition on molecular features rather than "
        "memorizing individual solutions.",
        align="justify")

    # QSCI scaling table
    qsci = (consolidated_phase3 or {}).get("sections", {}).get("qsci_scaling", {})
    mols = qsci.get("molecules", [])
    if mols:
        headers = ["Molecule", "Qubits", "Terms", "QSCI E (Ha)", "HF E (Ha)", "Backend"]
        rows = []
        for m in mols:
            rows.append([
                m["molecule"][:15],
                str(m["n_qubits"]),
                str(m["n_hamiltonian_terms"]),
                f"{m.get('qsci_energy', 0):.4f}" if m.get("qsci_energy") else "N/A",
                f"{m.get('hf_energy', 0):.4f}" if m.get("hf_energy") else "N/A",
                m.get("backend", "nvidia")[:10],
            ])
        add_table(doc, headers, rows)

    add_para(doc,
        "H2 (4q) QSCI recovers exact FCI energy. Benzene CAS(20e,20o) at 40 qubits completes "
        "in ~19 seconds on MPS with D=64. MPS bond dimension sweep shows stable energies "
        "across D=32-256, indicating the HF-dominated regime is well-captured by low-rank "
        "tensor networks. The 40-qubit result earns the GIC scaling bonus point.",
        align="justify")

    add_para(doc, "Figure 1. QWC circuit reduction", bold="")
    add_image(doc, FIG_DIR / "fig_qwc_reduction.png",
              "Figure 1: Qubit-wise commuting Pauli grouping reduces the number of circuits 2-3.5x, enabling cost-effective QPU execution.", width=5.0)

    add_heading(doc, "3.1 Identified Bottlenecks", level=2)
    add_para(doc,
        "We identify two concrete bottlenecks for scaling. First, the cuStateVec distributed "
        "statevector mode segfaults on PCIe-only L40S systems (no NVLink), capping exact "
        "simulation at 24 qubits. NVLink-equipped B200 systems remove this barrier, as "
        "demonstrated on qBraid's H200 instance. Second, diagonal sequence collapse: on "
        "strongly correlated systems (LiH, BeH2, N2 at stretched geometries), the model "
        "under-generates entangling X/Y operators and produces commuting Z-only sequences "
        "that are trapped at the Hartree-Fock energy. We traced this to the fixed-theta "
        "energy proxy used during RL sampling, which is nearly flat across candidate "
        "sequences and provides negligible gradient signal. Substituting a "
        "truncated L-BFGS-B reward directly targets this mechanism.",
        align="justify")

    add_para(doc, "Figure 2. Fixed-theta proxy vs converged energy", bold="")
    add_image(doc, FIG_DIR / "fig_02_proxy_flat_vs_final_varied.png",
              "Figure 2: The fixed-theta proxy is flat (top), while converged energies show real structure (bottom). This explains limited RL gains on large molecules.", width=5.5)

    # ===== PAGE 3: Accuracy & QPU Validation =====
    doc.add_page_break()
    add_heading(doc, "4. Accuracy and Benchmarking", level=1)

    bench = (consolidated_phase3 or {}).get("sections", {}).get("benchmark_ch3i", {})
    ref_E = bench.get("reference_energy", -6889.840354)
    methods = bench.get("methods", [])

    add_para(doc,
        f"We benchmark H-cGQE against Hardware-Efficient Ansatz VQE (HEA-VQE) and CUDA-Q GQE "
        f"on methyl iodide (CH3I) in a CAS(4e,4o) active space (8 qubits, 185 Hamiltonian "
        f"terms). Reference energy (CASCI/FCI): {ref_E:.6f} Ha.",
        align="justify")

    if methods:
        headers = ["Method", "Energy (Ha)", "Error (mHa)", "Runtime (s)"]
        rows = []
        for m in methods:
            rows.append([
                m["method"],
                f"{m['energy_hartree']:.6f}",
                f"{m['error_mha']:.3f}",
                f"{m.get('wall_time_seconds', 0):.1f}",
            ])
        add_table(doc, headers, rows)

    add_para(doc,
        "H-cGQE achieves 0.629 mHa error, outperforming both HEA-VQE (987.8 mHa, barren "
        "plateaus in the 8-qubit landscape) and CUDA-Q GQE (2.646 mHa, fixed operator pool "
        "without learned conditioning). On H2 (4q), QSCI recovers exact FCI energy (0.000 "
        "mHa). GPU benchmark across 17 molecules shows errors ranging from 0.0 mHa (H2 at "
        "equilibrium) to 817.6 mHa (N2 at 2.5 Angstrom), with 4 molecules reaching chemical "
        "accuracy (1.6 mHa).",
        align="justify")

    add_para(doc, "Figure 3. GPU benchmark: H-cGQE vs baselines", bold="")
    add_image(doc, FIG_DIR / "fig_gpu_benchmark_bar.png",
              "Figure 3: H-cGQE vs CUDA-Q GQE energy error relative to FCI across 17 molecules. H-cGQE reaches chemical accuracy on H2 and methyl iodide.", width=5.5)

    add_heading(doc, "4.1 QPU Validation on Rigetti Cepheus-1-108Q", level=2)

    qpu_results = (consolidated or {}).get("qpu_results", [])
    if qpu_results:
        add_para(doc,
            "We submit SQD sampling circuits for 12 molecules (8-28 qubits) to Rigetti "
            "Cepheus-1-108Q via qBraid with 8192 shots per job. Bitstrings are post-processed "
            "with SQD, including particle-number and spin-parity symmetry filtering. A "
            "critical bit-ordering fix was applied: Rigetti QPU bitstrings are reversed to "
            "match Qiskit convention (qubit 0 = LSB = rightmost) before SQD energy computation.",
            align="justify")

        headers = ["Molecule", "Qubits", "Electrons", "SQD Energy (Ha)", "FCI Energy (Ha)", "Error (mHa)"]
        rows = []
        for q in sorted(qpu_results, key=lambda x: x.get("n_qubits", 0)):
            mol = q.get("molecule", "?")[:18]
            nq = str(q.get("n_qubits", "?"))
            ne = str(q.get("n_electrons", "?"))
            sqd_e = q.get("sqd_energy", 0)
            fci_e = q.get("fci_energy")
            err = q.get("error_mHa")
            sqd_str = f"{sqd_e:.4f}" if sqd_e else "N/A"
            fci_str = f"{fci_e:.4f}" if fci_e else "N/A"
            err_str = f"{err:.1f}" if err is not None else "N/A"
            rows.append([mol, nq, ne, sqd_str, fci_str, err_str])
        add_table(doc, headers, rows)

        errors = [q.get("error_mHa") for q in qpu_results if q.get("error_mHa") is not None]
        best_err = min(errors) if errors else None
        max_qubits = max(q.get("n_qubits", 0) for q in qpu_results)
        add_para(doc,
            f"Across {len(qpu_results)} molecules on Cepheus-1-108Q (up to {max_qubits} "
            f"qubits), best SQD error vs FCI is {best_err:.1f} mHa (methyl iodide, 12q). "
            f"Seven EUV photoresist molecules achieve 13.9-53.3 mHa. We report "
            f"particle-number preservation (PNP) alongside every QPU energy: methyl iodide "
            f"attains chemical accuracy in noiseless simulation, and its 13.9 mHa hardware "
            f"figure tracks 8% PNP rather than a deficient circuit, isolating hardware noise "
            f"as the dominant error source. Ethylene (28q) and iodobenzene (8q) lack FCI "
            f"references for their active spaces; SQD energies are reported against HF where "
            f"available.",
            align="justify")

    add_para(doc, "Figure 4. QPU validation on Cepheus", bold="")
    add_image(doc, FIG_DIR / "fig_qpu_validation.png",
              "Figure 4: Rigetti Cepheus-1-108Q SQD results. Best QPU error is methyl iodide at 13.9 mHa. Lower PNP correlates with larger SQD error.", width=5.5)

    qpu_val = (consolidated_phase3 or {}).get("sections", {}).get("qpu_validation", {})
    if qpu_val and qpu_val.get("submissions"):
        add_para(doc,
            "Additionally, the CH3I H-cGQE circuit was validated on IQM Emerald 5-qubit QPU "
            "via AWS Braket, achieving 87.5% state fidelity (896/1024 shots in the expected "
            "state) with 4096 shots. qBraid simulator and AWS SV1 served as noise-free "
            "controls.")

    # ===== PAGE 4: Algorithmic Innovation & Hybrid System Design =====
    doc.add_page_break()
    add_heading(doc, "5. Algorithmic Innovation", level=1)

    add_para(doc, "H-cGQE advances beyond standard VQE and the original GQE in four areas:")

    add_para(doc,
        "(1) DAPO Reinforcement Learning: After supervised fine-tuning (96.2% validation "
        "accuracy), the model is refined with Decoupled Clip + Dynamic Sampling Policy "
        "Optimization. We sample operator sequences, evaluate energies on CUDA-Q's "
        "nvidia-mqpu target across 3 GPUs, and update the policy with asymmetric clipping "
        "(clip_low=0.2, clip_high=0.28) and GRPO-style group-relative advantages. Dynamic "
        "sampling skips flat-reward batches, and auxiliary rewards are gated on energy "
        "improvement over Hartree-Fock to prevent reward hacking.",
        align="justify")

    add_para(doc,
        "(2) MAP-Elites Quality-Diversity Archive: Elite circuits are binned by entanglement "
        "density and circuit depth, maintaining a diverse population that prevents mode "
        "collapse toward shallow, low-entanglement solutions. This archive directly addresses "
        "the diagonal sequence collapse identified in Section 3.1.",
        align="justify")

    add_para(doc,
        "(3) FMO2 Molecular Fragmentation: We implement Fragment Molecular Orbital many-body "
        "expansion to decompose iodobenzene into fragment subsystems. The FMO2 energy is "
        "reconstructed as E_FMO2 = sum(E_i) + sum(E_ij - E_i - E_j), where monomer and dimer "
        "energies are computed independently. This establishes correctness of the dimer "
        "correction; reducing maximum circuit width below the parent system requires three "
        "or more fragments.",
        align="justify")

    add_para(doc,
        "(4) SMILES Molecular Encoder: A chemistry-aware tokenizer (handling multi-character "
        "atoms like Cl, Br, Li, Be) feeds a 2-layer transformer (4 heads, 512 FFN dim, "
        "~202K parameters) to produce 256-dim molecular embeddings. Cosine similarity "
        "analysis shows chemically meaningful structure (N2-LiH 0.79, ethylene-CH3I 0.78), "
        "providing structural priors for conditioning on unseen molecular geometries.",
        align="justify")

    add_para(doc, "Figure 5. SFT vs RL ablation", bold="")
    add_image(doc, FIG_DIR / "fig_sft_vs_rl_ablation.png",
              "Figure 5: SFT-only vs SFT + DAPO RL. RL provides improvement on H2 but marginal changes on larger molecules due to the flat proxy reward.", width=5.5)

    add_heading(doc, "6. Hybrid System Design", level=1)
    add_para(doc,
        "The pipeline decouples HPC computation from QPU queue time. The HPC (AIRE L40S "
        "cluster or qBraid B200) performs all heavy computation: RL training, circuit "
        "synthesis, and L-BFGS-B optimization. It exports a QWC manifest containing grouped "
        "operators, optimized thetas, QASM circuits, and measurement bases. The manifest is "
        "submitted as a single batch to Rigetti Cepheus-1-108Q via qBraid, with async polling "
        "and retry logic (6 retries with exponential backoff for transient 404 errors). "
        "Retrieval and SQD post-processing happen separately, allowing the HPC to proceed "
        "with other work during QPU queue time.",
        align="justify")

    add_para(doc,
        "Error mitigation is integrated with hardware-aware preflight checks: REM (readout "
        "error correction via assignment matrix inversion) is applied when qubits <= 10; "
        "ZNE (zero-noise extrapolation via gate folding with scale factors [1,2,3] and "
        "Richardson extrapolation) is applied when two-qubit gates <= 20. SQD provides the "
        "primary noise resilience layer by filtering bitstrings through particle-number and "
        "spin-parity symmetry constraints before classical diagonalization in the selected "
        "determinant subspace.",
        align="justify")

    add_para(doc, "Figure 6. GPU scaling ladder", bold="")
    add_image(doc, FIG_DIR / "fig_gpu_scaling_ladder.png",
              "Figure 6: GPU scaling ladder from AIRE L40S (24q statevector) to qBraid B200 (32q) and 4xB200 (36q). NVLink removes the PCIe IPC bottleneck.", width=5.0)

    # ===== PAGE 5: Results Summary + Conclusion + Limitations =====
    doc.add_page_break()
    add_heading(doc, "7. Results Summary", level=1)

    headers = ["Experiment", "Metric", "Value", "Status"]
    rows = [
        ["QSCI Scaling", "Max qubits", "40", "Bonus point"],
        ["QSCI Benzene (40q)", "Runtime", "19.1 s", "MPS D=64"],
        ["QWC Grouping", "Circuit reduction", "2-3.5x", "H2: 15->5"],
        ["H-cGQE CH3I", "Error vs FCI", "0.629 mHa", "Chem. accuracy"],
        ["HEA-VQE CH3I", "Error vs FCI", "987.8 mHa", "Baseline"],
        ["CUDA-Q GQE CH3I", "Error vs FCI", "2.646 mHa", "Baseline"],
        ["QSCI H2 (4q)", "Error vs FCI", "0.000 mHa", "Exact"],
        ["Cepheus QPU SQD", "Molecules", "12 (8-28q)", "Rigetti 108Q"],
        ["Cepheus Best SQD", "Error vs FCI", "13.9 mHa", "Methyl iodide"],
        ["Cepheus EUV PR", "Molecules", "8 photoresist", "Mitsubishi Chem."],
        ["IQM Emerald QPU", "State fidelity", "87.5%", "CH3I circuit"],
        ["FMO2 Iodobenzene", "Solver error", "26.25 mHa", "Fragmentation"],
        ["SMILES Encoder", "Molecules", "10 (4-56q)", "Structural priors"],
        ["Error Mitigation", "Methods", "REM + ZNE + SQD", "Noise-aware"],
        ["Credit Usage", "Credits", "~11,475 / 13,400", "85.6% used"],
    ]
    add_table(doc, headers, rows)

    add_heading(doc, "8. Conclusion", level=1)
    add_para(doc,
        "We demonstrated a complete H-cGQE pipeline for the GIC Phase 3 competition: "
        "(1) QSCI scaling from 4 to 40 qubits (benzene CAS(20e,20o)) on MPS backend in under "
        "20 seconds, earning the GIC scaling bonus point. "
        "(2) QWC measurement grouping reduces QPU circuit count 2-3.5x, enabling batch "
        "submission within qBraid's 2000-circuit limit. "
        "(3) H-cGQE achieves 0.629 mHa on CH3I (chemical accuracy), outperforming HEA-VQE "
        "and CUDA-Q GQE. "
        "(4) 12 molecules (8-28 qubits) validated on Rigetti Cepheus-1-108Q with SQD "
        "post-processing, including 8 EUV photoresist molecules relevant to Mitsubishi "
        "Chemical. Best QPU SQD error: 13.9 mHa (methyl iodide, 12q). "
        "(5) DAPO RL with MAP-Elites archive directly optimizes for ground-state energy, "
        "with reward gating to prevent reward hacking. "
        "(6) A hybrid HPC-to-QPU workflow decouples classical computation from QPU queue "
        "time via async QWC manifest submission. "
        "(7) QPU submissions to Rigetti Cepheus and IQM Emerald via qBraid platform.",
        align="justify")

    add_heading(doc, "9. Limitations and Future Work", level=1)
    add_para(doc,
        "Strongly correlated systems remain the principal open challenge. While weakly correlated "
        "molecules near equilibrium reach chemical accuracy, multireference systems such as N2 at "
        "stretched geometries retain a 127 mHa gap to FCI. We traced this to the fixed-theta energy "
        "proxy used during RL sampling, which is nearly flat across candidate sequences (Fig. 2) and "
        "therefore supplies little gradient signal for non-trivial circuit topologies. Substituting a "
        "truncated L-BFGS-B reward and symmetry-preserving constrained decoding directly targets this "
        "mechanism and is the clear next step.",
        align="justify")
    add_para(doc,
        "On hardware, SQD error grows with qubit count and circuit depth, from 13.9 mHa (methyl iodide, "
        "12q) to 130.0 mHa (N2, 20q), consistent with per-gate two-qubit fidelity near 99.1% on Cepheus. "
        "We report particle-number preservation (8-71%) alongside every QPU energy precisely so that "
        "hardware noise can be separated from ansatz quality: methyl iodide attains chemical accuracy "
        "in noiseless simulation, and its 13.9 mHa hardware figure tracks 8% particle-number "
        "preservation rather than a deficient circuit. Deeper mitigation and hardware-aware "
        "transpilation are the natural levers.",
        align="justify")
    add_para(doc,
        "Two results are deliberately scoped. QSCI at 40 qubits demonstrates that the MPS pipeline "
        "executes at that width; the sampled subspace is HF-dominated, so the recovered energy is "
        "Hartree-Fock rather than correlated, and post-HF accuracy at this scale requires deeper "
        "entangling circuits. FMO2 fragmentation is validated on a two-fragment partition of "
        "iodobenzene, establishing correctness of the dimer correction; reducing maximum circuit width "
        "below the parent system requires three or more fragments. Neither affects the benchmark or "
        "QPU results reported above.",
        align="justify")

    # ===== REFERENCES (not counted in page limit) =====
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    refs = [
        "[1] Kanno et al., Phys. Rev. A 108, 022405 (2023). GQE for ground-state energy.",
        "[2] Robledo-Moreno et al., Nature 634, 795-800 (2024). Chemistry beyond exact solutions on a quantum computer.",
        "[3] Kanno et al., arXiv:2409.03657 (2024). Generative quantum eigensolver with transformer.",
        "[4] Yu et al., arXiv:2503.14476 (2025). DAPO: Open-source LLM RL system.",
        "[5] Mouret & Clune, arXiv:1504.04909 (2015). Illuminating search spaces by mapping elites.",
        "[6] Temme et al., Nature 567, 209-212 (2019). Error mitigation with ZNE.",
        "[7] Bravyi et al., arXiv:2003.04997 (2020). REM for readout error correction.",
        "[8] Peruzzo et al., Nat. Commun. 5, 4213 (2014). VQE on a photonic quantum processor.",
        "[9] Grimsley et al., Nat. Commun. 10, 3007 (2019). ADAPT-VQE.",
        "[10] NVIDIA CUDA-Q, https://nvidia.github.io/cuda-quantum/",
        "[11] qBraid platform, https://www.qbraid.com/",
        "[12] Fedorov & Kitamura, Int. J. Quantum Chem. 107, 2227 (2007). FMO method.",
    ]
    for r in refs:
        add_para(doc, r, size=10)

    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOC))
    print(f"Saved: {OUT_DOC}")


if __name__ == "__main__":
    build_document()
